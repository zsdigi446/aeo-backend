"""
Word 报告生成器 v2
适配新报告结构（AEO Expert 方法论），按 lang 输出双语 Word 文档
"""

import io
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from analyzer.i18n_report import WORD_I18N


def _w(lang: str, key: str) -> str:
    return WORD_I18N.get(lang, WORD_I18N["zh-CN"])[key]


def generate_word(report: dict, lang: str = "zh-CN") -> bytes:
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    DARK_BLUE = RGBColor(0x17, 0x36, 0x5D)
    meta = report.get("meta", {})

    # ===== 标题 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{meta.get('site_name', '')} {_w(lang, 'report_title_suffix')}")
    run.font.size = Pt(26)
    run.font.color.rgb = DARK_BLUE
    run.bold = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    pPr = title._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="4" w:color="4F81BD"/></w:pBdr>')
    pPr.append(pBdr)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(_w(lang, 'subtitle'))
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"{_w(lang, 'analysis_target')}{meta.get('url', '')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    score_p = doc.add_paragraph()
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = score_p.add_run(f"{_w(lang, 'prelim_score')} {meta.get('total_score', 0)} / 100")
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_BLUE
    run.bold = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph()

    # ===== Part 1: 核心结论 =====
    p1 = report.get("part1_core_judgment", {})
    _add_h1(doc, p1.get("title", ""))

    _add_body(doc, p1.get("overview_score", ""), bold=True, size=12)
    _add_body(doc, p1.get("summary", ""), size=11)
    doc.add_paragraph()
    _add_body(doc, p1.get("judgment", ""), size=11)
    doc.add_paragraph()

    if p1.get("dimension_summary"):
        for ds in p1["dimension_summary"]:
            _add_body(doc, f"• {ds}", size=11)

    doc.add_paragraph()
    _add_body(doc, p1.get("priority_action", ""), bold=True, size=11)

    # 评分表格
    dimensions = p1.get("dimensions", [])
    if dimensions:
        doc.add_paragraph()
        table = doc.add_table(rows=len(dimensions) + 1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = [_w(lang, 'th_dimension'), _w(lang, 'th_weight'), _w(lang, 'th_score'), _w(lang, 'th_key_finding')]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="365F91" w:val="clear"/>')
            cell._element.get_or_add_tcPr().append(shading)

        for r, dim in enumerate(dimensions):
            data = [dim.get("name", ""), f"{dim.get('weight', 0)}%", f"{dim.get('score', 0)}/100", dim.get("key_finding", "")]
            for c, val in enumerate(data):
                cell = table.rows[r + 1].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(9)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                if r % 2 == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
                    cell._element.get_or_add_tcPr().append(shading)

    # ===== Part 2: AEO 优势 =====
    _add_h1(doc, report.get("part2_advantages", {}).get("title", ""))
    for item in report.get("part2_advantages", {}).get("items", []):
        p = doc.add_paragraph()
        run = p.add_run(f"• {item}")
        run.font.size = Pt(11)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.paragraph_format.left_indent = Cm(0.5)

    # ===== Part 3: AEO 问题 =====
    _add_h1(doc, report.get("part3_problems", {}).get("title", ""))
    for prob in report.get("part3_problems", {}).get("problems", []):
        _add_h2(doc, prob.get("title", ""))
        _add_body(doc, prob.get("detail", ""), size=11)

    # ===== Part 4: 内容类型覆盖 =====
    p4 = report.get("part4_content_coverage", {})
    if p4:
        _add_h1(doc, p4.get("title", ""))
        _add_body(doc, p4.get("description", ""), size=11)
        doc.add_paragraph()
        for ct in p4.get("content_types", []):
            icon = "✅" if ct["covered"] else "❌"
            _add_body(doc, f"{icon} {ct['type']}（优先级：{ct['priority']}）— {ct['description']}", size=10)

    # ===== Part 5: Persona × Funnel × Use Case =====
    p5 = report.get("part5_opportunities", {})
    _add_h1(doc, p5.get("title", ""))
    _add_body(doc, p5.get("description", ""), size=11)

    scenarios = p5.get("scenarios", [])
    if scenarios:
        doc.add_paragraph()
        table = doc.add_table(rows=len(scenarios) + 1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = [_w(lang, 'sc_persona'), _w(lang, 'sc_funnel'), _w(lang, 'sc_use_case'), _w(lang, 'sc_ai_question'), _w(lang, 'sc_recommended_page')]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="365F91" w:val="clear"/>')
            cell._element.get_or_add_tcPr().append(shading)

        for r, sc in enumerate(scenarios):
            data = [sc.get("persona", ""), sc.get("funnel", ""), sc.get("use_case", ""), sc.get("ai_question", ""), sc.get("page", "")]
            for c, val in enumerate(data):
                cell = table.rows[r + 1].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(8)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ===== Part 6: 优先页面 =====
    p6 = report.get("part6_priority_pages", {})
    _add_h1(doc, p6.get("title", ""))
    _add_body(doc, p6.get("description", ""), size=11)
    for group in p6.get("groups", []):
        _add_h2(doc, group.get("name", ""))
        for i, page in enumerate(group.get("pages", []), 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {page}")
            run.font.size = Pt(11)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            p.paragraph_format.left_indent = Cm(0.5)

    top5 = p6.get("top5", [])
    if top5:
        p = doc.add_paragraph()
        run = p.add_run(f"{_w(lang, 'top5_label')}{_w(lang, 'sep').join(top5[:5])}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ===== Part 7: 页面模板 =====
    p7 = report.get("part7_page_template", {})
    _add_h1(doc, p7.get("title", ""))
    _add_body(doc, p7.get("description", ""), size=11)

    example = p7.get("example", {})
    p = doc.add_paragraph()
    run = p.add_run(f"{_w(lang, 'example_page_label')}{example.get('page_title', '')}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    for item in example.get("structure", []):
        level = item["level"]
        if level == "H1":
            _add_h2(doc, item["content"])
        elif level == "H2":
            p = doc.add_paragraph()
            run = p.add_run(f"• {item['content']}")
            run.font.size = Pt(11)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            p.paragraph_format.left_indent = Cm(0.5)
        elif level in ("对比表", "Comparison Table"):
            p = doc.add_paragraph()
            run = p.add_run(f"📊 {item['content']}")
            run.font.size = Pt(11)
            run.italic = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        else:
            p = doc.add_paragraph()
            run = p.add_run(item["content"])
            run.font.size = Pt(11)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    _add_h2(doc, _w(lang, 'geo_template_label'))
    _add_body(doc, example.get("geo_template", ""), size=11)

    # 八要素清单
    if example.get("eight_elements"):
        doc.add_paragraph()
        _add_body(doc, _w(lang, 'eight_elements_label'), bold=True, size=11)
        for el in example["eight_elements"]:
            _add_body(doc, f"  {el}", size=10)

    # ===== Part 8: 技术建议 =====
    p8 = report.get("part8_technical", {})
    _add_h1(doc, p8.get("title", ""))
    for item in p8.get("items", []):
        p = doc.add_paragraph()
        run = p.add_run(f"• {item}")
        run.font.size = Pt(11)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.paragraph_format.left_indent = Cm(0.5)

    # ===== Part 9: 效果衡量 =====
    p9 = report.get("part9_measurement", {})
    _add_h1(doc, p9.get("title", ""))
    _add_body(doc, p9.get("description", ""), size=11)

    for dim in p9.get("dimensions", []):
        _add_h2(doc, dim.get("name", ""))
        _add_body(doc, dim.get("description", ""), size=11)
        for prompt in dim.get("prompts", []):
            p = doc.add_paragraph()
            run = p.add_run(f"  • {prompt}")
            run.font.size = Pt(10)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x36, 0x5F, 0x91)
    run.bold = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def _add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)
    run.bold = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def _add_body(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
